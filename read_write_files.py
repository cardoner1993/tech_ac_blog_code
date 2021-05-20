import pandas as pd
from docx import Document
from tqdm import tqdm
from PIL import Image, ImageFile
import requests

ImageFile.LOAD_TRUNCATED_IMAGES = True


def add_image(url, img_name):
    im = Image.open(requests.get(url, stream=True).raw)
    im = im.resize((500, 600), Image.ANTIALIAS)
    im1 = im.save(img_name, "PNG")
    return img_name


dfs = pd.read_excel('/home/daca/Downloads/ATP Daily - 1st - 2021_05_11.xlsx', sheet_name='Sheet1', skiprows=4,
                    header=1)

vars_to_use = ['Photolink', 'Wholesale Price EUR', 'SKU', 'Style Color', 'Style Description', 'Color Name Long ID']

document = Document()
document.add_heading('Bolsos Kipling 2021_05_11', 0)

p = document.add_paragraph()
r = p.add_run()

try:
    for id, row in tqdm(dfs[vars_to_use].iterrows(), total=len(dfs)):

        if requests.get(row['Photolink'], stream=True).raw.status == 200:
            r.add_picture(add_image(row['Photolink'], img_name=f"{row['SKU']}.png"))
            r.add_text(f"Modelo "
                       f"{', '.join(row[['SKU', 'Style Color', 'Style Description', 'Color Name Long ID']].values)}\n")
            r.add_text(f"Precio en EUR {row['Wholesale Price EUR']}\n")
            r.add_text("#############################################")

except KeyboardInterrupt as e:
    print(f"Error. Reason {str(e)}")

finally:
    document.save('bolsos_Kipling_2021_05_11.docx')